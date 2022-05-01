'''
main report classes
'''

import os
from docx import Document
import pandas as pd

from docx2pdf import convert
from datetime import  datetime
from reports.AQ_Data import  AQ_Data, get_data_files
import reports.Limits as Limits
import reports.breathe_safety_index_api as breathe_safety_index_api
from pathlib import Path
from reports.utils import replace_dict_key, delete_empty_column, set_color
import pytz
from reports.HVAC_param_translation import HVAC_KEY_Translation
from reports.params import ( HOME_MONTH_TEMPLATE, HOME_DAY_TEMPLATE, 
                    OFFICE_DAY_TEMPLATE, OFFICE_MONTH_TEMPLATE, templates_path, DOCX_HOME_REPORTS_PATH, DOCX_OFFICE_REPORTS_PATH)  

productivity_ordering = {
        "Integral" : 1,
        "Basic activity": 2,
        "Applied activity": 3,
        "Focused activity": 4,
        "Task orientation": 5,
        "Initiative": 6,
        "Information search": 7,
        "Information usage": 8,
        "Breadth of approach": 9,
        "Basic strategy": 10,
    }

life_quality_ordering = {
    "Mental health": 1,
    "Sleep quality": 2,
    "Conditions for physical activity" : 3,
}

health_risks_ordering = {
    "CVS": 1,
    "Brain": 2,
    "Respiratory" : 3,
    "Sensorial" : 4,
    "Allergy": 5
}

body_systems = {
    "CVS": "cardiovascular system",
    "Brain": "brain health (nervous system)",
    "Respiratory" : "respiratory",
    "Sensorial" :  "sensorial health",
    "Allergy": "allergy (immune system)"
}
HVAC_ordering = {
    "temperature": 1,
    "humidity": 2,
    "CO2": 3,
    "lux": 4,
    "dB_average": 5,
    "PM2_5": 6,
    "PM10": 7,
    "CH2O": 8,
    "O3": 9,
}

class PeriodMixing():
    def get_period_name(self) -> str:
        pass

    def fill_periods_in_table(self, table):
        pass

    def get_period_str(self) -> str:
        pass

    def months_or_days(self) -> str:
        pass       

class DayMixing(PeriodMixing):
    
    def get_period_name(self) -> str:
        return 'hours'

    def fill_periods_in_table(self, table):
        for i in range(self.period_len):
            table.cell(i+1,0).text = str(self.aq_periods["time"].iloc[i])

    def get_period_str(self) -> str:
        return f'{str(self.day)}-{str(self.month)}-{str(self.year)}'

    def months_or_days(self) -> str:
        return 'DAYS'

class MonthMixing(PeriodMixing):

    def get_period_name(self) -> str:
        return "days"

    def fill_periods_in_table(self, table):

         for i in range(self.period_len):
            table.cell(i+1,0).text = str(self.aq_periods["date_day"].iloc[i])

    def get_period_str(self) -> str:
        start = str(self.aq_periods["date_day"].iloc[0])
        end = str(self.aq_periods["date_day"].iloc[len(self.aq_periods)-1])
        period_str = f"{start}: {end}"
        return period_str

    def months_or_days(self) -> str:
        return 'MONTHS'

class Report():
    
    def __init__(self, month: int, year: int, device: str, aq_data: AQ_Data, aq_periods: pd.DataFrame, template_name: str, fahrenheit = False) -> None:
        self.fahrenheit = fahrenheit
        self.month = month
        self.year = year
        self.device = device
        self.aq_data = aq_data
        self.doc_report = Document(template_name)
        self.aq_periods = aq_periods
        self.period_len = len(self.aq_periods)

    def set_report_var_value(self, vars: dict) -> None:
        paragraphs = self.doc_report.paragraphs
        for paragraph in paragraphs:
            for name in vars.keys():
                if name in paragraph.text:
                    paragraph.text = paragraph.text.replace(name, vars[name])

    def add_rows_to_table(self, table):
        for i in range(self.period_len):
            if i < self.period_len - 1:
                table.add_row()
          
    def get_HVAC_table(self):
        return self.doc_report.tables[0]
    
    def get_health_risks_table(self):
        return self.doc_report.tables[2]

    def fill_period_and_device(self):
        
        period_str = self.get_period_str()

        self.set_report_var_value({"{{period}}": period_str})
        self.set_report_var_value({"{{device}}": self.device})       
            
    def fill_HVAC_table(self):
        table = self.get_HVAC_table()
        self.add_rows_to_table(table)
        
        for i in range(self.period_len):
                        
            temperature_color = Limits.temperature("temperature",self.aq_periods["temperature"].iloc[i], self.fahrenheit)
            set_color(table.cell(i+1,1), temperature_color)
            table.cell(i+1,1).text = str(round(self.aq_periods["temperature"].iloc[i]))

            for key in Limits.HVAC_color_functions:
                if key != "temperature":
                    if key in self.aq_periods.columns:
                        hvac_param = HVAC_KEY_Translation[key]
                        color = Limits.HVAC_color_functions[key](hvac_param, self.aq_periods[key].iloc[i])
                        column = HVAC_ordering[key]
                        set_color(table.cell(i+1, column), color)
                        try:
                            value = str(round(self.aq_periods[key].iloc[i]))
                        except:
                            value = "0.0"
                        table.cell(i+1,column).text = value

        delete_empty_column(table)

        self.fill_HVAC_conclusion()

    def fill_HVAC_conclusion(self):
                
        high_temp_aq_list = self.aq_periods[self.aq_periods["temperature"]>Limits.T_MAX ]
        high_temp_list = high_temp_aq_list["temperature"].tolist()  
        temp_high_count = len(high_temp_list)

        low_temp_aq_list = self.aq_periods[self.aq_periods["temperature"] < Limits.T_MIN ]
        low_temp_list = low_temp_aq_list["temperature"].tolist()
        temp_low_count = len(low_temp_list)
        
        if temp_high_count > 0:
            avgTempUp = round(sum(high_temp_list)/temp_high_count - Limits.T_MAX,1)

        if temp_low_count > 0:
            avgTempLess = round(Limits.T_MIN - sum(low_temp_list)/temp_low_count,1)

        humidity_high_count = len(self.aq_periods[self.aq_periods["humidity"]> Limits.HUMIDITY_HIGH])
        humidity_low_count = len(self.aq_periods[self.aq_periods["humidity"]< Limits.HUMIDITY_LOW])
        
        co2_high_count = len(self.aq_periods[self.aq_periods["CO2"]> Limits.CO2_LIMIT])
        dust_high_count = len(self.aq_periods[self.aq_periods["PM2_5"]> Limits.PM_2_5_LIMIT])
        
        periods_qty = len(self.aq_periods)

        hvac_conclusion_text = ""

        if temp_high_count > 0:
            if len(hvac_conclusion_text)>0:
                hvac_conclusion_text = f'{hvac_conclusion_text}\r\n'
            hvac_conclusion_text = (f'{hvac_conclusion_text}Temperature was higher {temp_high_count} {self.get_period_name()} out of  {periods_qty}.'
                                    f'You need to decrease heating level by {avgTempUp} degrees,'
                                    'that may decrease your bills (1 degree = 5% economy)'
                                    )

        if temp_low_count > 0:
            if len(hvac_conclusion_text)>0:
                hvac_conclusion_text = f'{hvac_conclusion_text}\r\n'
            hvac_conclusion_text = (f'{hvac_conclusion_text}Temperature was lower {temp_low_count} {self.get_period_name()} out of {periods_qty}, '
                                    f'you need to increase heating level by {avgTempLess} degrees, '
                                    'that may increase your bills up to (1 degree = 5% increase) but improve  workplace environment and job satisfaction'
                                    )

        if humidity_high_count > 0:
            if len(hvac_conclusion_text)>0:
                hvac_conclusion_text = f'{hvac_conclusion_text}\r\n'
            hvac_conclusion_text = (f'{hvac_conclusion_text}Humidity was higher {humidity_high_count} {self.get_period_name()} out of  {periods_qty}, '
                                    'you need to install/adjust a dehumidifier or ventilation'
                                    )
        if humidity_low_count > 0:
            if len(hvac_conclusion_text)>0:
                hvac_conclusion_text = f'{hvac_conclusion_text}\r\n'
            hvac_conclusion_text = (f'{hvac_conclusion_text}Humidity was lower {humidity_low_count} {self.get_period_name()} out of  {periods_qty}, '
                                    'you need to install/adjust a humidifier'
                                    )

        if co2_high_count > 0:
            if len(hvac_conclusion_text)>0:
                hvac_conclusion_text = f'{hvac_conclusion_text}\r\n'
            hvac_conclusion_text = (f'{hvac_conclusion_text}CO2 was higher {co2_high_count} {self.get_period_name()} out of {periods_qty}, '
                                    'you need to improve ventilation rate at your place'
                                    )
        if dust_high_count > 0:
            if len(hvac_conclusion_text)>0:
                hvac_conclusion_text = f'{hvac_conclusion_text}\r\n'
            hvac_conclusion_text = (f'{hvac_conclusion_text}Dust was higher {dust_high_count} {self.get_period_name()} out of {periods_qty},'
                                    'you need to improve ventilation rate at room, install a dust filter or service it.'
                                    )

        if hvac_conclusion_text == "":
            hvac_conclusion_text = "All air quality parameters are within recommended ranges"

        self.set_report_var_value ({"{{HVAC efficiency conclusion}}": hvac_conclusion_text})
           
    def aq_data_prepare(self):
        period_dict = self.aq_periods.to_dict('index')
        self.aq_data_list = []
        for key in period_dict.keys():
            try:
                period_dict[key].pop("date_day")
                period_dict[key].pop("time")
                period_dict[key].pop("date")
            except:
                pass
      
            data_dict = period_dict[key]

            for key in HVAC_KEY_Translation:
                try:
                    replace_dict_key(data_dict, key, HVAC_KEY_Translation[key])
                except:
                    pass
            
            self.aq_data_list.append(data_dict)
    
    def get_all_risks(self):
        risk_list = breathe_safety_index_api.get_all_risks(self.aq_data_list)
        
        self.health_risks = []
        self.productivity = []
        self.life_quality = []
        
        self.productivity_impacts = []
        self.health_risks_impacts = []
        self.life_quality_impacts = []

        for risk in risk_list:
            health_risk = {}
            productivity = {}
            life_quality = {}
            
            for body_system in list(health_risks_ordering.keys()):
                health_risk[body_system] = risk[body_system]['risk']
                #health_risk[body_system] = risk[0][body_system]['risk']
            for direction in list(productivity_ordering.keys()):
                if direction == "Integral":
                    productivity[direction] = risk['Productivity'][direction]
                else:
                    productivity[direction] = risk['Productivity'][direction]['value']

            life_quality['Sleep quality'] = risk['SleepQuality']['risk']
            life_quality['Conditions for physical activity'] = risk['Conditions']['risk']
            life_quality['Mental health'] = risk['Mental']['risk']

            self.health_risks.append(health_risk)
            self.productivity.append(productivity)
            self.life_quality.append(life_quality)
    
    """ 
    def get_productivity_impacts(self, risk_list):
        self.productivity_impacts = {}

    def get_health_risks_impacts(self, risk_list):
        self.health_risks_impacts = {}
        
    def get_life_quality_impacts(self, risk_list):
        self.life_quality_impacts = {}
    """

    def fill_health_risks_table(self):
        table = self.get_health_risks_table()
        self.add_rows_to_table(table)
        
        for i in range(self.period_len):
            
            for key in health_risks_ordering.keys():
                color = Limits.health_risks(self.health_risks[i][key])
                column = health_risks_ordering[key]
                set_color(table.cell(i+1, column), color)
                try:
                    value = str(round(self.health_risks[i][key],1))
                except:
                    value = "0.0"
                table.cell(i+1,column).text = value
        
        self.fill_health_risks_conclusion()

    def fill_health_risks_conclusion(self):

        risky_systems = {}
        list_health_risks = {k: [dic[k] for dic in self.health_risks] for k in self.health_risks[0]}

        for system in health_risks_ordering.keys():
            try:
                if max(list_health_risks[system]) >=3:
                    try:
                        risky_systems[system] += 1
                    except:
                        risky_systems[system] = 1
            except:
                pass

        if len(risky_systems.keys()) > 0:

            system_list = (list(risky_systems.keys()))
            body_system_list = []
            for system in system_list:
                body_system_list.append (body_systems[system])
                
            systems_str = ', '.join(body_system_list)
            health_risks_conclusion_text = f"Risk of short-term labour incapacity (2-3 days) for {systems_str}"
        else:
            health_risks_conclusion_text = "Direct disease risks are not identified or are implicit."

        self.set_report_var_value ({"{{Health risks conclusion}}": health_risks_conclusion_text})    
      
    def fill_periods_in_tables(self):
        tables = self.doc_report.tables
        tables.pop(1) # delete table with health risks score
        for table in tables:
            self.fill_periods_in_table(table)

    def fill_report(self):

        self.fill_period_and_device()
        
        self.fill_HVAC_table()
        self.aq_data_prepare()
        self.get_all_risks()
        self.fill_health_risks_table()
        
    def fill_device(self):
        pass

class OfficeReport(Report):   

    def __init__(self, month: int, year: int, device: str, aq_data: AQ_Data, aq_periods: pd.DataFrame, template_name: str, fahrenheit = False) -> None:
        super().__init__(month, year, device, aq_data, aq_periods=aq_periods, template_name = template_name, fahrenheit= fahrenheit)
    
    def get_productivity_table(self):
        return self.doc_report.tables[3]

    def fill_productivity_table(self):
        table = self.get_productivity_table()
        self.add_rows_to_table(table)
        
        keys = productivity_ordering.keys()
        
        for i in range(self.period_len):
            
            for key in keys:
                value = self.productivity[i][key]
                color = Limits.productivity(value)
                column = productivity_ordering[key]
                set_color(table.cell(i+1, column), color)
                table.cell(i+1, column).text = str (value)+ "%"
        
        self.fill_productivity_conclusion()

    def fill_productivity_conclusion(self):
        productivity_conclusion_text = ''
        risky_directions = []
        safe_directions = []
        directions = list(productivity_ordering.keys())
        directions.remove("Integral")
        for direction in directions:
            direction_values = list(map(lambda productivity: productivity[direction], self.productivity))
            direction_value_avg = sum(direction_values) / len(direction_values)
            if direction_value_avg < 80:
                risky_directions.append(direction)
            else:
                safe_directions.append(direction)

        if len(safe_directions) > 0:
            safe_directions_text = f'Values for loss of productivity are within the normal range for {", ".join(safe_directions)}'
        else:
            safe_directions_text = ""

        if len(risky_directions) > 0:
            risky_directions_text = f'Values for loss of productivity are in risky zone for {", ".join(risky_directions)}'
        else:
            risky_directions_text = ""

        if safe_directions_text != "":
            productivity_conclusion_text = f"{safe_directions_text} \r\n {risky_directions_text}"
        else:
            productivity_conclusion_text = risky_directions_text

        self.set_report_var_value({"{{Productivity conclusion}}": productivity_conclusion_text})
    
    def fill_report(self):
        super().fill_report()
                
        self.fill_productivity_table()
        self.fill_periods_in_tables()

        Path(self.path).mkdir(parents=True, exist_ok=True)

        self.doc_report.save(os.path.join(self.path, self.report_name))   

class HomeReport(Report):
    
    def __init__(self, month: int, year: int, device: str, aq_data: AQ_Data, aq_periods: pd.DataFrame, template_name: str, fahrenheit = False) -> None:
        super().__init__(month, year, device, aq_data, aq_periods=aq_periods, template_name = template_name, fahrenheit= fahrenheit)

    def get_life_quality_table(self):
        return self.doc_report.tables[3]

    def fill_life_quality_table(self):
        table = self.get_life_quality_table()
        self.add_rows_to_table(table)
                
        for i in range(self.period_len):
                        
            for key in self.life_quality[i].keys():
                color = Limits.mental_health_color_functions[key](self.life_quality[i][key])
        
                column = life_quality_ordering[key]
                set_color(table.cell(i+1, column), color)
                table.cell(i+1, column).text = str (round(self.life_quality[i][key], 1)) # str(round(self.health_risks[i][key],1))
        
        self.fill_life_quality_conclusion()
    
    def fill_life_quality_conclusion(self):
        life_quality_conclusion_text = ''
        risky_directions = []
        safe_directions = []
        
        directions = list(life_quality_ordering.keys())
        for direction in directions:
            direction_values = list(map(lambda life_quality: life_quality[direction], self.life_quality))
            direction_avg = sum(direction_values)/len(direction_values)
            if direction == "Mental health":
                if direction_avg >= 3:
                    risky_directions.append(direction)
                else:
                    safe_directions.append(direction)
            elif direction_avg < 80:
                 risky_directions.append(direction)
            else:
                safe_directions.append(direction)

        if len(safe_directions) > 0:
            safe_directions_text = f'Values for life quality are within the normal range for {", ".join(safe_directions)}'
        else:
            safe_directions_text = ""

        if len(risky_directions) > 0:
            risky_directions_text = f'Values for life quality are in risky zone for {", ".join(risky_directions)}'
        else:
            risky_directions_text = ""

        if safe_directions_text != "":
            life_quality_conclusion_text = f"{safe_directions_text} \r\n {risky_directions_text}"
        else:
            life_quality_conclusion_text = risky_directions_text

        self.set_report_var_value({"{{Life quality conclusion}}": life_quality_conclusion_text})

    def fill_report(self):

        super().fill_report()
        
        self.fill_life_quality_table()
        self.fill_periods_in_tables()

        Path(self.path).mkdir(parents=True, exist_ok=True)
        self.doc_report.save(os.path.join(self.path, self.report_name))
        
class HomeMonthReport(HomeReport, MonthMixing):
    def __init__(self, month, year, device : str, aq_data: AQ_Data) -> None:

        aq_periods = aq_data.get_month_by_nights(month = month, year = year)
              
        template = os.path.join(templates_path, HOME_MONTH_TEMPLATE)
        self.report_name = f'{device}_{month}_{year}_home_report.docx'
        
        super().__init__(month, year, device, aq_periods=aq_periods, template_name = template, aq_data = aq_data)
        self.path = os.path.join(DOCX_HOME_REPORTS_PATH, self.months_or_days(), self.device, str(self.year))

class HomeDayReport(HomeReport, DayMixing):
    def __init__(self, month, year, day, device : str, aq_data: AQ_Data) -> None:
        
        self.day = day
        
        aq_periods = aq_data.get_night_data_by_hours(year = year, month = month, day = day)
        
        template = os.path.join(templates_path, HOME_DAY_TEMPLATE)
        
        self.report_name = f'{device}_{day}_home_report.docx'
        
        super().__init__(month, year, device, aq_periods=aq_periods, template_name = template, aq_data = aq_data)
        self.path = os.path.join(DOCX_HOME_REPORTS_PATH, self.months_or_days(), self.device, str(self.year))
        self.path = os.path.join(self.path, str(self.month))      

class OfficeDayReport(OfficeReport, DayMixing):
    def __init__(self, month: int, year: int, day: int, device : str, aq_data: AQ_Data) -> None:
        
        self.day = day
        
        aq_periods = aq_data.get_day_data_by_hours(year = year, month = month, day = day)
      
        template = os.path.join(templates_path, OFFICE_DAY_TEMPLATE)
        
        self.report_name = f'{device}_{day}_office_report.docx'
        
        super().__init__(month, year, device, aq_periods=aq_periods, template_name = template, aq_data = aq_data)
        self.path = os.path.join(DOCX_OFFICE_REPORTS_PATH, self.months_or_days(), self.device, str(self.year))
        self.path = os.path.join(self.path, str(self.month))
          
class OfficeMonthReport(OfficeReport, MonthMixing):
    
    def __init__(self, month, year, device : str, aq_data: AQ_Data) -> None:
        
        aq_periods = aq_data.get_month_by_days(month = month, year = year)
                
        template = os.path.join(templates_path, OFFICE_MONTH_TEMPLATE)
        self.report_name = f'{device}_{month}_{year}_office_report.docx'
       
        super().__init__(month, year, device, aq_periods=aq_periods, template_name = template, aq_data = aq_data)
        self.path = os.path.join(DOCX_OFFICE_REPORTS_PATH, self.months_or_days(), self.device, str(self.year))
                                          
    def export_report(self):
        self.doc_report =os.path.join(self.path, self.report_name)
        self.pdf_name = self.report_name.replace('docx', 'pdf')
        convert (self.report_name, self.pdf_name)

def day_reports_from_file(aq_folder, aq_file, report_type: str, tz: pytz.timezone):

    aq_data = AQ_Data(aq_folder, aq_file, tz)
    
    for _, month in enumerate(aq_data.months):
        print("         Current Time =", datetime.now().strftime("%H:%M:%S"))
        print(month)
        days = aq_data.get_month_dates(year = month[1], month = month[0])
        for day in days:
            print(day)
            if report_type.upper() =="HOME":
                r =  HomeDayReport(month = month[0], year = month[1], day = day, device =aq_file[:-4], aq_data = aq_data)
            elif report_type.upper() == "OFFICE":
                r =  OfficeDayReport(month = month[0], year = month[1], day = day, device =aq_file[:-4], aq_data = aq_data)
            else:
                return
        
            if not r.aq_data.empty:
                r.fill_report()

def month_reports_from_file(aq_folder, aq_file: str, report_type: str, tz: pytz.timezone):
    aq_data = AQ_Data(aq_folder, aq_file, tz)
   
    for _, month in enumerate(aq_data.months):
        
        print("Current Time =", datetime.now().strftime("%H:%M:%S"))
        print(month)
        if report_type.upper() =="HOME":
            r =  HomeMonthReport(month = month[0], year = month[1], device =aq_file[:-4], aq_data = aq_data)
        elif report_type.upper() == "OFFICE":
            r =  OfficeMonthReport(month = month[0], year = month[1], device =aq_file[:-4], aq_data = aq_data)
        else:
            return
        if r.period_len != 0:
            r.fill_report()

def home_day_reports_from_file(aq_folder, aq_file ):
    pass

def office_day_reports_from_file(aq_folder, aq_file ):
    pass

def test_day_report(aq_folder, aq_file):
    aq_data = AQ_Data(aq_folder, aq_file)
    
    r = HomeDayReport(year = 2021, month = 9, day = 2, aq_data = aq_data, device = aq_file[:-4])
    r.fill_report()
    
if __name__ == "__main__":
   
    aq_files, aq_folders = get_data_files()

    #folder = aq_folders[0]
    #aq_file = aq_files[folder][0]

    t_start = datetime.now()

    #day_reports_from_file(folder, aq_file, "Home")

    
    """ for folder in aq_folders:
        print(folder)
        for aq_file in aq_files[folder]:
            print(aq_file)
            month_reports_from_file(folder, aq_file, "HOME")
            day_reports_from_file(folder, aq_file, "HOME")
            month_reports_from_file(folder, aq_file, "OFFICE")
            day_reports_from_file(folder, aq_file, "OFFICE")
 """
    folder = aq_folders[1] 
    for aq_file in aq_files[folder][5:]:
        print(aq_file)
        if "JP" in folder:
            tz = pytz.timezone("Japan")
        elif "US" in folder:
            tz = pytz.timezone("US/Pacific")
        month_reports_from_file(folder, aq_file, "HOME", tz)
        day_reports_from_file(folder, aq_file, "HOME", tz)
        month_reports_from_file(folder, aq_file, "OFFICE", tz)
        day_reports_from_file(folder, aq_file, "OFFICE", tz)

    t_end = datetime.now()
    print(t_end-t_start)